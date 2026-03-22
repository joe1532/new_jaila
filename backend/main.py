import base64
import io
import json
import logging
import os
import re
import time
import unicodedata
from pathlib import Path
from uuid import uuid4

import fitz
import openpyxl
import xlrd
from docx import Document
from fastapi import FastAPI, File, Header, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from openai import OpenAI
from pypdf import PdfReader

from backend.config import (
    BASE_DIR,
    CHAT_INSTRUCTIONS,
    LOG_DIR,
    PRIMARY_MODEL,
    PROMPT_CACHE_KEY_CHAT,
    PROMPT_CACHE_KEY_LIGNINGSFRIST,
    PROMPT_CACHE_RETENTION,
    REASONING_EFFORT_CHAT,
    REASONING_EFFORT_LIGNINGSFRIST,
    SAGSBEHANDLING_MODELS,
    SAGSBEHANDLING_PROMPTS,
    SAGSBEHANDLING_VECTOR_STORES,
    get_allowed_origins,
)
from backend.models import (
    AnalyzeRequest,
    AnalyseLogGetResponse,
    AnalyseLogListResponse,
    AnalyseLogSaveRequest,
    AnalyseLogSaveResponse,
    AnalyzeResponse,
    CaseCreateRequest,
    CaseGetResponse,
    CaseListResponse,
    CaseUpdateRequest,
    ChatContextFileResponse,
    ChatContextListResponse,
    ChatExportRequest,
    ChatExportResponse,
    ChatLogGetResponse,
    ChatLogListResponse,
    ChatLogSaveRequest,
    ChatLogSaveResponse,
    ChatRequest,
    ChatResponse,
    LegalSourcesCatalogResponse,
    LegalSourceSectionResponse,
    SagsLegalBasisResponse,
)
from backend.services.analyse_logs import (
    delete_analyse_log,
    format_log_as_text,
    get_analyse_log,
    list_analyse_logs,
    save_analyse_log,
)
from backend.services.chat_logs import delete_chat_log, get_chat_log, list_chat_logs, save_chat_log
from backend.services.case_store import get_case_store
from backend.services.openai_service import analyze_question, analyze_question_stream
from backend.services.pdf_log import save_chat_pdf_log, save_pdf_log


app = FastAPI(title="JAILA Backend API", version="1.0.0")
_log = logging.getLogger(__name__)
CHAT_CONTEXT_DIR = BASE_DIR / "chat_context"
DEFAULT_CHAT_GUIDE_PATH = BASE_DIR / "backend" / "default_context" / "Skriveguide.md"
MAX_CHAT_CONTEXT_CHARS = 20000
MAX_CHAT_CONTEXT_PER_FILE_CHARS = 30000
MAX_UPLOAD_BYTES = 15 * 1024 * 1024
MAX_PDF_PAGES = 20
MAX_PDF_OCR_PAGES = 8
MIN_PDF_PAGE_TEXT_FOR_NO_OCR = 40
MAX_DOCX_PARAGRAPHS = 1200
MAX_EXCEL_SHEETS = 10
MAX_EXCEL_ROWS_PER_SHEET = 400
MAX_SAGS_CONTEXT_CHARS = 6000
MAX_SAGS_CONTEXT_LOGS = 8
MAX_ANALYSE_LEGAL_CONTEXT_BLOCKS = 10
MAX_ANALYSE_LEGAL_CONTEXT_CHARS = 12000
LEGAL_SOURCES_DIR = Path(os.getenv("LEGAL_SOURCES_DIR", "/var/lib/jaila/legal_sources")).resolve()
LEGAL_SOURCE_PREVIEW_CACHE: dict[str, dict[str, object]] = {}
LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE: dict[str, dict[str, dict[str, object]]] = {}
LEGAL_SOURCE_PRECOMPUTED_MTIME_BY_NAMESPACE: dict[str, float] = {}
LEGAL_PREVIEW_REMOVE_PATTERNS = [
    re.compile(r"printet fra karnov til brug i overensstemmelse med licensvilk[aå]rene", re.IGNORECASE),
]
LEGAL_PREVIEW_DEFAULT_START_MARKER_PATTERN = re.compile(
    r"er\s+blevet\s+enige\s+om\s+f[oø]lgende\s*:?",
    re.IGNORECASE,
)
LEGAL_PREVIEW_TYSKLAND_START_MARKER_PATTERN = re.compile(
    r"der\s+har\s+til\s+hensigt\s*:?",
    re.IGNORECASE,
)

allowed_origins = get_allowed_origins()
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)


@app.get("/api/health")
def health() -> dict[str, str]:
    if not os.getenv("OPENAI_API_KEY"):
        return {"status": "degraded", "reason": "OPENAI_API_KEY mangler"}
    return {"status": "ok"}


def sanitize_filename(filename: str) -> str:
    cleaned = re.sub(r"[^\w.\-]+", "_", filename.strip(), flags=re.UNICODE)
    cleaned = cleaned.strip("._")
    return cleaned or "context.txt"


def normalize_text(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = cleaned.strip()
    return cleaned


def normalize_match_text(text: str) -> str:
    lowered = str(text or "").lower()
    normalized = unicodedata.normalize("NFKD", lowered)
    without_marks = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", without_marks).strip()


def normalize_wrapped_preview_lines(text: str) -> str:
    """Join hard line wraps from PDF extraction without removing paragraph breaks."""
    value = str(text or "")
    # Join hyphenated line breaks, e.g. "for-\nhold" -> "forhold".
    value = re.sub(r"(?<=[A-Za-zÆØÅæøå0-9])-\n(?=[A-Za-zÆØÅæøå0-9])", "", value)
    # Join single line breaks where next line starts lowercase (typical hard-wrap mid-sentence).
    value = re.sub(r"(?<=[A-Za-zÆØÅæøå0-9,;:)\]])\n(?=[a-zæøå])", " ", value)
    return value


def clean_legal_preview_text(text: str, namespace: str = "") -> str:
    """Remove repeated license/footer noise from preview text."""
    lines = [line.strip() for line in str(text or "").splitlines()]
    kept: list[str] = []
    for line in lines:
        if not line:
            kept.append("")
            continue
        if any(pattern.search(line) for pattern in LEGAL_PREVIEW_REMOVE_PATTERNS):
            continue
        kept.append(line)
    normalized = "\n".join(kept)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = normalized.strip()
    namespace_key = str(namespace or "").strip().lower()
    marker_match = None
    if namespace_key == "tyskland":
        marker_match = LEGAL_PREVIEW_TYSKLAND_START_MARKER_PATTERN.search(normalized)
    if not marker_match:
        marker_match = LEGAL_PREVIEW_DEFAULT_START_MARKER_PATTERN.search(normalized)
    if marker_match:
        normalized = normalized[marker_match.end() :].strip()
    normalized = normalize_wrapped_preview_lines(normalized).strip()
    return normalized


def parse_legal_source_namespace(source_id: str) -> tuple[str, str]:
    safe_id = str(source_id or "").strip().lower()
    if not re.fullmatch(r"[a-z0-9_]{5,160}", safe_id):
        raise HTTPException(status_code=400, detail="Ugyldigt source_id-format")
    namespace = safe_id.split("_dbo_", 1)[0] if "_dbo_" in safe_id else ""
    if not namespace:
        raise HTTPException(status_code=400, detail="Kunne ikke udlede namespace fra source_id")
    return safe_id, namespace


def resolve_legal_pdf_by_source_id(source_id: str) -> Path:
    safe_id, namespace = parse_legal_source_namespace(source_id)
    legal_files_dir = (LEGAL_SOURCES_DIR / namespace / "files").resolve()
    if not legal_files_dir.exists():
        raise HTTPException(status_code=404, detail="Retskilde-filer er ikke tilgængelige")

    article_match = re.search(r"_art(\d{1,2})_v\d{4}$", safe_id)
    is_protocol = bool(re.search(r"_protokol_v\d{4}$", safe_id))
    document_slug_match = re.search(r"_dbo_(.+?)_v\d{4}$", safe_id)
    document_slug = str(document_slug_match.group(1) if document_slug_match else "").strip().lower()
    normalized_target = None
    if article_match:
        normalized_target = f"artikel {int(article_match.group(1))}"
    elif is_protocol:
        normalized_target = "protokol"

    candidates = sorted(path for path in legal_files_dir.glob("*.pdf") if path.is_file())
    for candidate in candidates:
        stem_norm = normalize_match_text(candidate.stem)
        stem_slug = re.sub(r"[^a-z0-9]+", "_", stem_norm).strip("_")
        if is_protocol and "protokol" in stem_norm:
            return candidate
        if normalized_target and normalized_target in stem_norm:
            return candidate
        if document_slug and document_slug == stem_slug:
            return candidate

    raise HTTPException(status_code=404, detail="PDF-kilde blev ikke fundet for source_id")


def extract_pdf_preview_text(file_path: Path, namespace: str = "", max_chars: int = 14000) -> tuple[str, bool]:
    """Extract plain text preview from PDF for in-app reading."""
    try:
        reader = PdfReader(str(file_path))
        text_parts: list[str] = []
        for page in reader.pages:
            chunk = normalize_text(page.extract_text() or "")
            if chunk:
                text_parts.append(chunk)
            if sum(len(part) for part in text_parts) >= max_chars:
                break
        joined = "\n\n".join(text_parts).strip()
        if not joined:
            return "Kunne ikke udtrække tekst fra PDF.", False
        cleaned = clean_legal_preview_text(joined, namespace=namespace)
        if not cleaned:
            return "Kunne ikke udtrække tekst fra PDF.", False
        return truncate_text(cleaned, max_chars)
    except Exception:
        return "Kunne ikke udtrække tekst fra PDF.", False


def extract_pdf_preview_pages(file_path: Path, namespace: str = "") -> list[str]:
    """Extract and clean text per page for paginated preview."""
    try:
        reader = PdfReader(str(file_path))
        pages: list[str] = []
        for page in reader.pages:
            chunk = clean_legal_preview_text(normalize_text(page.extract_text() or ""), namespace=namespace)
            if chunk:
                pages.append(chunk)
        return pages
    except Exception:
        return []


def load_precomputed_legal_previews(namespace: str) -> dict[str, dict[str, object]]:
    preview_path = (LEGAL_SOURCES_DIR / f"{namespace}_previews.json").resolve()
    if not preview_path.exists():
        LEGAL_SOURCE_PRECOMPUTED_MTIME_BY_NAMESPACE[namespace] = -1.0
        LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE[namespace] = {}
        return LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE[namespace]
    try:
        current_mtime = float(preview_path.stat().st_mtime)
    except Exception:
        current_mtime = -1.0
    current_cache = LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE.get(namespace)
    current_cache_mtime = float(LEGAL_SOURCE_PRECOMPUTED_MTIME_BY_NAMESPACE.get(namespace, -2.0))
    if current_cache is not None and current_cache_mtime == current_mtime:
        return current_cache
    try:
        payload = json.loads(preview_path.read_text(encoding="utf-8"))
        entries = payload.get("entries") if isinstance(payload, dict) else {}
        if isinstance(entries, dict):
            LEGAL_SOURCE_PRECOMPUTED_MTIME_BY_NAMESPACE[namespace] = current_mtime
            LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE[namespace] = entries
            return entries
    except Exception:
        pass
    LEGAL_SOURCE_PRECOMPUTED_MTIME_BY_NAMESPACE[namespace] = current_mtime
    LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE[namespace] = {}
    return LEGAL_SOURCE_PRECOMPUTED_CACHE_BY_NAMESPACE[namespace]


def truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars] + "\n\n[NOTE: Indhold afkortet pga. længde.]", True


def strip_pdf_suffix(filename: str) -> str:
    return re.sub(r"\.pdf$", "", str(filename or "").strip(), flags=re.IGNORECASE).strip()


def list_vector_store_document_names(client: OpenAI, vector_store_id: str) -> list[str]:
    names: list[str] = []
    seen: set[str] = set()
    after: str | None = None

    # NOTE: Keep pagination defensive because SDK response shape may vary by version.
    for _ in range(10):
        kwargs = {"vector_store_id": vector_store_id, "limit": 100}
        if after:
            kwargs["after"] = after
        page = client.vector_stores.files.list(**kwargs)
        data = getattr(page, "data", None) or []
        if not data:
            break

        for item in data:
            item_filename = str(getattr(item, "filename", "") or "")
            file_id = str(getattr(item, "file_id", "") or "")
            item_id = str(getattr(item, "id", "") or "")

            candidate_name = item_filename
            if not candidate_name:
                openai_file_id = file_id or (item_id if item_id.startswith("file-") else "")
                if openai_file_id:
                    try:
                        file_obj = client.files.retrieve(openai_file_id)
                        candidate_name = str(getattr(file_obj, "filename", "") or "")
                    except Exception:
                        candidate_name = ""

            cleaned = strip_pdf_suffix(candidate_name)
            if not cleaned:
                continue
            key = cleaned.casefold()
            if key in seen:
                continue
            seen.add(key)
            names.append(cleaned)

        has_more = bool(getattr(page, "has_more", False))
        if not has_more:
            break
        last_item = data[-1]
        after = str(getattr(last_item, "id", "") or "")
        if not after:
            break

    return sorted(names, key=lambda value: value.casefold())


def normalize_fact_value(value: object) -> str:
    text = str(value or "").strip()
    return re.sub(r"\s+", " ", text).strip()


def format_case_facts_for_llm(subtab: str, case_facts: dict[str, object] | None) -> str:
    facts = case_facts or {}
    if subtab != "skattepligt_ligningsfrist":
        return ""

    mapping: list[tuple[str, str]] = [
        ("income_years", "Indkomstår"),
        ("foreign_income", "Forhold der kan begrunde ordinær ligningsfrist"),
        ("foreign_assets_liabilities", "Aktiver/passiver i udlandet"),
        ("residence_fact", "Bopælsfaktum"),
    ]

    lines: list[str] = []
    for key, label in mapping:
        value = normalize_fact_value(facts.get(key, ""))
        if not value:
            continue
        lines.append(f"- {label}: {value}")

    if not lines:
        return ""

    return (
        "Faktiske oplysninger fra sagsbehandler (skal indgå i vurderingen):\n"
        + "\n".join(lines)
    )


def build_sags_context_from_analyse_log(username: str, context_log_id: str) -> str:
    """Byg deterministisk sagskontekst fra én tidligere analyse-log."""
    safe_user = str(username or "").strip()
    safe_log_id = str(context_log_id or "").strip()
    if not safe_user or not safe_log_id:
        return ""

    entry = get_analyse_log(safe_user, safe_log_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Valgt analyse-kontekst blev ikke fundet")

    answer_text = str(entry.get("answer", "") or "").strip()
    if not answer_text:
        raise HTTPException(status_code=400, detail="Valgt analyse-log har intet svar at bruge som kontekst")

    title = str(entry.get("title", "Uden titel") or "Uden titel").strip()
    created_at = str(entry.get("created_at", "") or "").strip()
    question = str(entry.get("question", "") or "").strip()
    context_text = (
        "Tidligere analyse-kontekst (gennemgået af bruger)\n"
        f"- Titel: {title}\n"
        f"- Tidspunkt: {created_at}\n"
        f"- Spørgsmål: {question}\n"
        "Tidligere analyse/svar:\n"
        f"{answer_text}"
    )
    context_text, _ = truncate_text(context_text, MAX_SAGS_CONTEXT_CHARS)
    return context_text


def build_sags_context_from_analyse_logs(username: str, context_log_ids: list[str]) -> str:
    """Byg samlet sagskontekst fra flere analyse-logs."""
    safe_user = str(username or "").strip()
    if not safe_user or not context_log_ids:
        return ""

    blocks: list[str] = []
    total_chars = 0
    chars_per_log = max(500, MAX_SAGS_CONTEXT_CHARS // len(context_log_ids))

    for log_id in context_log_ids:
        safe_log_id = str(log_id or "").strip()
        if not safe_log_id:
            continue
        block = build_sags_context_from_analyse_log(safe_user, safe_log_id)
        if not block:
            continue
        block, _ = truncate_text(block, chars_per_log)
        blocks.append(block)
        total_chars += len(block)
        if total_chars >= MAX_SAGS_CONTEXT_CHARS:
            break

    if not blocks:
        return ""
    combined = "\n\n---\n\n".join(blocks)
    combined, _ = truncate_text(combined, MAX_SAGS_CONTEXT_CHARS)
    return combined


def format_case_shared_facts_for_llm(case_entry: dict | None, active_subtab: str) -> str:
    """Lav kompakt, deterministisk kontekstblok fra sag på tværs af undertabs."""
    if not case_entry:
        return ""
    shared_facts = case_entry.get("shared_facts") or {}
    subtab_outputs = case_entry.get("subtab_outputs") or {}
    lines: list[str] = []
    if isinstance(shared_facts, dict) and shared_facts:
        lines.append("Fælles sagsfakta (tidligere fastlagt i sagen):")
        for key, value in shared_facts.items():
            if value is None:
                continue
            text = str(value).strip()
            if not text:
                continue
            lines.append(f"- {key}: {text}")
    if isinstance(subtab_outputs, dict):
        for subtab, output in subtab_outputs.items():
            if not isinstance(output, dict):
                continue
            if str(subtab or "").strip() == str(active_subtab or "").strip():
                continue
            answer = str(output.get("answer", "") or "").strip()
            if not answer:
                continue
            lines.append(f"Tidligere delresultat fra undertab '{subtab}':")
            lines.append(answer[:2000])
    if not lines:
        return ""
    block = "\n".join(lines)
    block, _ = truncate_text(block, MAX_SAGS_CONTEXT_CHARS)
    return block


def format_sags_decision_package_for_llm(decision_package: object | None) -> str:
    """Lav kompakt, deterministisk tekstblok fra struktureret beslutningspakke."""
    if not decision_package:
        return ""

    if hasattr(decision_package, "model_dump"):
        try:
            package = decision_package.model_dump()
        except Exception:
            package = {}
    elif isinstance(decision_package, dict):
        package = decision_package
    else:
        package = {}

    if not isinstance(package, dict) or not package:
        return ""

    lines: list[str] = ["Beslutningspakke (struktureret vurderingsgrundlag):"]
    sagskontekst = package.get("sagskontekst") or {}
    if isinstance(sagskontekst, dict):
        indkomsttype = normalize_fact_value(sagskontekst.get("indkomsttype"))
        bopaelsland = normalize_fact_value(sagskontekst.get("bopaelsland"))
        arbejdsgivertype = normalize_fact_value(sagskontekst.get("arbejdsgivertype"))
        if indkomsttype:
            lines.append(f"- Sagskontekst / indkomsttype: {indkomsttype}")
        if bopaelsland:
            lines.append(f"- Sagskontekst / bopælsland: {bopaelsland}")
        if arbejdsgivertype:
            lines.append(f"- Sagskontekst / arbejdsgivertype: {arbejdsgivertype}")
        selected_article = sagskontekst.get("valgt_artikel") or {}
        if isinstance(selected_article, dict):
            article = selected_article.get("article")
            section = selected_article.get("section")
            raw_text = normalize_fact_value(selected_article.get("raw_text"))
            if article:
                article_label = f"artikel {article}"
                if section:
                    article_label += f", stk. {section}"
                lines.append(f"- Sagskontekst / valgt artikel: {article_label}")
            elif raw_text:
                lines.append(f"- Sagskontekst / valgt artikel (rå): {raw_text}")

    rule_profile = package.get("regelprofil") or {}
    if isinstance(rule_profile, dict):
        profile_id = normalize_fact_value(rule_profile.get("profile_id"))
        if profile_id:
            lines.append(f"- Regelprofil: {profile_id}")

    constated_facts = package.get("konstaterede_fakta") or []
    if isinstance(constated_facts, list) and constated_facts:
        lines.append("Konstaterede fakta:")
        for fact in constated_facts[:20]:
            if not isinstance(fact, dict):
                continue
            fact_key = normalize_fact_value(fact.get("fact_key"))
            value = fact.get("value")
            value_text = normalize_fact_value(value if isinstance(value, str) else json.dumps(value, ensure_ascii=False))
            if fact_key and value_text:
                lines.append(f"- {fact_key}: {value_text}")

    premises = package.get("afledte_praemisser") or []
    if isinstance(premises, list) and premises:
        lines.append("Afledte præmisser:")
        for premise in premises[:12]:
            text = normalize_fact_value(premise)
            if text:
                lines.append(f"- {text}")

    legal_sources = package.get("relevante_retskilder") or []
    if isinstance(legal_sources, list) and legal_sources:
        lines.append("Relevante retskilder:")
        for source in legal_sources[:12]:
            if isinstance(source, dict):
                label = normalize_fact_value(source.get("label") or source.get("title") or source.get("name"))
                reason = normalize_fact_value(source.get("reason") or source.get("why"))
                if label and reason:
                    lines.append(f"- {label} ({reason})")
                elif label:
                    lines.append(f"- {label}")
            else:
                text = normalize_fact_value(source)
                if text:
                    lines.append(f"- {text}")

    unresolved = package.get("uafklarede_sporgsmaal") or []
    if isinstance(unresolved, list) and unresolved:
        lines.append("Uafklarede spørgsmål:")
        for item in unresolved[:12]:
            text = normalize_fact_value(item)
            if text:
                lines.append(f"- {text}")

    allocation = package.get("fordelingsmetode") or {}
    if isinstance(allocation, dict):
        method_id = normalize_fact_value(allocation.get("method_id"))
        description = normalize_fact_value(allocation.get("description"))
        if method_id or description:
            lines.append("Fordelingsmetode:")
            if method_id:
                lines.append(f"- Metode: {method_id}")
            if description:
                lines.append(f"- Beskrivelse: {description}")

    preliminary_tax_right = package.get("foreloebig_beskatningsret") or []
    if isinstance(preliminary_tax_right, list) and preliminary_tax_right:
        lines.append("Foreløbig beskatningsret:")
        for item in preliminary_tax_right[:20]:
            if not isinstance(item, dict):
                continue
            country = normalize_fact_value(item.get("country"))
            label = normalize_fact_value(item.get("label"))
            basis = normalize_fact_value(item.get("basis"))
            if country or label or basis:
                parts = [part for part in [country, label, basis] if part]
                lines.append(f"- {' | '.join(parts)}")

    conflicts = package.get("konflikter") or []
    if isinstance(conflicts, list) and conflicts:
        lines.append("Konflikter:")
        for item in conflicts[:12]:
            text = normalize_fact_value(item)
            if text:
                lines.append(f"- {text}")

    warnings = package.get("advarsler") or []
    if isinstance(warnings, list) and warnings:
        lines.append("Advarsler:")
        for item in warnings[:12]:
            text = normalize_fact_value(item)
            if text:
                lines.append(f"- {text}")

    if len(lines) <= 1:
        return ""
    block = "\n".join(lines)
    block, _ = truncate_text(block, MAX_SAGS_CONTEXT_CHARS)
    return block


def parse_income_years(value: object) -> list[int]:
    raw = str(value or "").strip()
    years = re.findall(r"\b((?:19|20)\d{2})\b", raw)
    if not years:
        raise HTTPException(status_code=400, detail="Angiv mindst ét indkomstår (YYYY)")
    unique_sorted_years = sorted({int(year) for year in years})
    return unique_sorted_years


def calculate_ordinær_frist_year(income_year: int) -> int:
    return income_year + 4


def get_short_deadline_regulations(income_years: list[int]) -> list[str]:
    regulations: list[str] = []
    if any(year <= 2023 for year in income_years):
        regulations.append("1305 af 14. november 2018")
    if any(year >= 2024 for year in income_years):
        regulations.append("49 af 24. januar 2025")
    return regulations


def format_danish_list(values: list[str]) -> str:
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    if len(values) == 2:
        return f"{values[0]} og {values[1]}"
    return f"{', '.join(values[:-1])} og {values[-1]}"


def format_income_years_display(income_years: list[int]) -> str:
    """Årstal til indkomstperiode: 1 år = '2023', 2 år = '2022 og 2023', 3+ år = '2022-2024'."""
    year_texts = [str(y) for y in income_years]
    if not year_texts:
        return ""
    if len(year_texts) == 1:
        return year_texts[0]
    if len(year_texts) == 2:
        return f"{year_texts[0]} og {year_texts[1]}"
    return f"{year_texts[0]}-{year_texts[-1]}"


def format_income_years_label(income_years: list[int]) -> str:
    years_display = format_income_years_display(income_years)
    if not years_display:
        return ""
    if len(income_years) == 1:
        return f"For indkomståret {years_display}"
    return f"For indkomstårene {years_display}"


def format_deadline_lines(income_years: list[int]) -> str:
    if len(income_years) == 1:
        income_year = income_years[0]
        frist_year = calculate_ordinær_frist_year(income_year)
        return (
            f"For indkomståret {income_year} kan vi derfor varsle ændring senest den 1. maj {frist_year} "
            f"og foretage ansættelsen senest den 1. august {frist_year}."
        )

    lines: list[str] = []
    lines.append(f"{format_income_years_label(income_years)} kan vi derfor varsle og ansætte således:")
    for year in income_years:
        frist_year = calculate_ordinær_frist_year(year)
        lines.append(
            f"- Indkomståret {year}: varsling senest den 1. maj {frist_year} og ansættelse senest den 1. august {frist_year}."
        )
    return "\n".join(lines)


def select_single_trigger(case_facts: dict[str, object]) -> str:
    selected_trigger = str(case_facts.get("selected_trigger", "")).strip()
    selected_factors_raw = case_facts.get("selected_factors")
    selected_factors: list[str] = []
    if isinstance(selected_factors_raw, list):
        for value in selected_factors_raw:
            normalized = str(value or "").strip()
            if normalized:
                selected_factors.append(normalized)

    candidates = {selected_trigger} if selected_trigger else set()
    candidates.update(selected_factors)
    if len(candidates) != 1:
        raise HTTPException(status_code=400, detail="Vælg præcis én trigger")
    return next(iter(candidates))


def build_residence_clause(case_facts: dict[str, object]) -> str:
    mode = str(case_facts.get("residence_mode", "")).strip()
    since_year_raw = str(case_facts.get("residence_since_year", "")).strip()

    if mode == "always":
        return "Da du altid har haft bopæl i Danmark"

    if mode == "since_year":
        if not re.search(r"\b(?:19|20)\d{2}\b", since_year_raw):
            raise HTTPException(
                status_code=400,
                detail="Angiv gyldigt årstal for bopæl i Danmark siden",
            )
        return f"Da du har haft bopæl i Danmark siden {since_year_raw}"

    # Backward compatibility for existing payloads with free text.
    legacy_text = str(case_facts.get("residence_fact", "")).strip()
    if not legacy_text:
        return ""
    legacy_text = legacy_text.rstrip(". ")
    if re.match(r"^\s*da\s+du\b", legacy_text, flags=re.IGNORECASE):
        return legacy_text
    legacy_text = re.sub(r"^\s*du\s+", "", legacy_text, flags=re.IGNORECASE)
    return f"Da du {legacy_text.strip()}"


def get_trigger_detail(case_facts: dict[str, object], trigger_id: str) -> str:
    selected_detail = str(case_facts.get("selected_trigger_detail", "")).strip()
    if selected_detail:
        return selected_detail
    detail_by_trigger_key = f"{trigger_id}_detail"
    return str(case_facts.get(detail_by_trigger_key, "")).strip()


def build_trigger_text(
    trigger_id: str,
    trigger_detail: str,
    self_employed_mode: str,
) -> tuple[str, bool, str, str, bool]:
    detail = str(trigger_detail or "").strip()
    mode_map = {
        "not_covered_by_section_2": "oplysningsskema",
        "with_annual_statement_exception_rule": "undtagelse",
    }
    normalized_self_employed_mode = mode_map.get(self_employed_mode, self_employed_mode)

    if trigger_id == "self_employed_business":
        # Ingen tekstboks længere – brug standardformulering når detail er tom
        effective_detail = detail or "enkeltmandsvirksomhed eller deltagelse i interessentskab"
        if normalized_self_employed_mode == "oplysningsskema":
            return (
                f"Da du har haft selvstændig erhvervsvirksomhed i form af {effective_detail}",
                True,
                "§ 1, stk. 2, nr. 1",
                "1",
                False,
            )
        if normalized_self_employed_mode == "undtagelse":
            return (
                f"Da du har haft selvstændig erhvervsvirksomhed i form af {effective_detail}",
                True,
                "§ 2",
                "2",
                False,
            )
        raise HTTPException(
            status_code=400,
            detail="Ugyldig underkategori for selvstændig erhvervsvirksomhed",
        )
    # Juridisk mapping (v1, deterministisk):
    # - uses_1302=True: self_employed_business, foreign_income, foreign_real_estate,
    #   work_abroad_with_relief, special_tax_liability_conditions
    # - uses_1302=False: major_shareholder_status,
    #   foreign_assets_liabilities_significant, cross_border_commuter_taxation
    trigger_map: dict[str, tuple[str, bool, str, str, bool]] = {
        "foreign_income": (
            f"Da du har haft indkomst fra udlandet i form af {detail}",
            True,
            "§ 3",
            "2",
            True,
        ),
        "foreign_real_estate": (
            "Da du har haft fast ejendom i udlandet",
            True,
            "§ 3",
            "2",
            False,
        ),
        "work_abroad_with_relief": (
            "Da du har haft lønindkomst for arbejde udført i udlandet, hvor der gives nedslag i dansk skat efter en dobbeltbeskatningsoverenskomst eller efter danske regler",
            True,
            "§ 1, stk. 2, nr. 5",
            "1",
            False,
        ),
        "special_tax_liability_conditions": (
            f"Da du har haft {detail}",
            True,
            "§ 1, stk. 3",
            "1",
            True,
        ),
        "major_shareholder_status": (
            f"Da du har haft hovedaktionærstatus i selskabet {detail}",
            False,
            "",
            "3",
            True,
        ),
        "foreign_assets_liabilities_significant": (
            f"Da du har haft formueforhold i udlandet af betydning for skatteansættelsen i form af {detail}",
            False,
            "",
            "4",
            True,
        ),
        "cross_border_commuter_taxation": (
            "Da du har været omfattet af grænsegængerbeskatning efter kildeskattelovens § 5 A til § 5 D",
            False,
            "",
            "5",
            False,
        ),
    }
    if trigger_id not in trigger_map:
        raise HTTPException(status_code=400, detail="Ugyldig trigger valgt")
    return trigger_map[trigger_id]


def build_standard_text_ligningsfrist(case_facts: dict[str, object] | None) -> str:
    facts = case_facts or {}
    income_years = parse_income_years(facts.get("income_years", ""))
    selected_trigger = select_single_trigger(facts)
    trigger_detail = get_trigger_detail(facts, selected_trigger)
    self_employed_mode = str(facts.get("self_employed_business_mode", "")).strip()
    if selected_trigger == "self_employed_business" and not self_employed_mode:
        raise HTTPException(
            status_code=400,
            detail="Vælg underkategori for selvstændig erhvervsvirksomhed",
        )
    trigger_text, uses_1302, relevant_bestemmelse, relevant_nummer, requires_detail = build_trigger_text(
        selected_trigger,
        trigger_detail,
        self_employed_mode,
    )
    if requires_detail and not trigger_detail:
        raise HTTPException(status_code=400, detail="Triggeren kræver supplerende tekst")

    bopæl_clause = build_residence_clause(facts) if selected_trigger != "cross_border_commuter_taxation" else ""
    if selected_trigger != "cross_border_commuter_taxation" and not bopæl_clause:
        raise HTTPException(status_code=400, detail="Bopælsfaktum skal udfyldes")

    regulation_labels = get_short_deadline_regulations(income_years)
    regulation_label_text = " og ".join(
        [f"bekendtgørelse nr. {label}" for label in regulation_labels]
    )
    income_years_text = format_income_years_display(income_years)
    income_period_label = (
        f"indkomståret {income_years_text}"
        if len(income_years) == 1
        else f"indkomstårene {income_years_text}"
    )
    deadline_lines = format_deadline_lines(income_years)

    if selected_trigger == "foreign_income":
        trigger_text = f"Da du har haft {trigger_detail} fra udlandet i {income_period_label}"
    elif selected_trigger == "foreign_real_estate":
        trigger_text = f"Da du har haft fast ejendom fra udlandet i {income_period_label}"
    elif selected_trigger == "work_abroad_with_relief":
        trigger_text = (
            "Da du har haft lønindkomst for arbejde udført i udlandet, "
            "hvor der gives nedslag i dansk skat efter en dobbeltbeskatningsoverenskomst "
            f"eller efter danske regler i {income_period_label}"
        )
    elif selected_trigger == "major_shareholder_status":
        trigger_text = f"Da du har haft hovedaktionærstatus i selskabet {trigger_detail} i {income_period_label}"
    elif selected_trigger == "foreign_assets_liabilities_significant":
        trigger_text = (
            f"Da du har haft formueforhold i udlandet som vi vurderer har betydning for "
            f"skatteansættelsen i form af {trigger_detail} i {income_period_label}"
        )
    elif selected_trigger == "special_tax_liability_conditions":
        special_mode = str(facts.get("special_tax_liability_mode", "")).strip()
        years_special = format_income_years_display(income_years)
        period_special = (
            f"indkomståret {years_special}"
            if len(income_years) == 1
            else f"indkomstårene {years_special}"
        ) if years_special else income_period_label
        if special_mode == "shift_full_limited":
            trigger_text = (
                f"Da du både har været fuldt og begrænset skattepligtig til Danmark i {period_special}"
            )
        elif special_mode == "tax_resident_abroad":
            trigger_text = (
                f"Da du har været skattemæssigt hjemmehørende i udlandet i {period_special}"
            )
        elif special_mode == "offset_income_year":
            trigger_text = f"Da du har forskudt indkomstår i {years_special}"
        elif special_mode == "duty_under_section_8_2":
            trigger_text = (
                f"Da du har haft oplysningspligt efter skattekontrollovens § 8, stk. 2 i {period_special}"
            )
        elif special_mode == "request_information_schema":
            trigger_text = (
                f"Da du har anmodet om oplysningsskema i {period_special}"
            )
        else:
            trigger_text = f"Da du har haft {trigger_detail} i {period_special}"

    if uses_1302:
        first_paragraph = (
            f"{trigger_text}, er det vores vurdering, at du er omfattet af {relevant_bestemmelse} i bekendtgørelse nr. 1302 af 14. november 2018 om fysiske personers modtagelse af en årsopgørelse i stedet for et oplysningsskema."
        )
        second_paragraph = (
            f"Efter § 2, stk. 1, nr. {relevant_nummer}, i {regulation_label_text} om en kort frist for skatteansættelse af personer med enkle økonomiske forhold anses du derfor ikke for at have enkle økonomiske forhold."
        )
    else:
        first_paragraph = (
            f"{trigger_text}, er det vores vurdering, at du er omfattet af § 2, stk. 1, nr. {relevant_nummer} i {regulation_label_text} om en kort frist for skatteansættelse af personer med enkle økonomiske forhold."
        )
        second_paragraph = "Du anses derfor ikke for at have enkle økonomiske forhold."

    if selected_trigger == "cross_border_commuter_taxation":
        return (
            "Med hjemmel i skatteforvaltningslovens § 26, stk. 1, har skatteministeren fastsat en kortere ligningsfrist for fysiske personer med enkle økonomiske forhold."
            + "\n\n"
            + first_paragraph
            + "\n\n"
            + second_paragraph
            + "\n\n"
            "Den korte ligningsfrist finder derfor ikke anvendelse, og du er dermed omfattet af den ordinære ligningsfrist i skatteforvaltningslovens § 26, stk. 1.\n\n"
            + deadline_lines
            + "\n\n"
            "Det fremgår af kildeskattelovens § 2, stk. 1, at personer, der ikke er fuldt skattepligtige til Danmark, men som erhverver indkomst fra arbejde udført her i landet, er begrænset skattepligtige til Danmark af denne indkomst.\n\n"
            "Personer, der er begrænset skattepligtige efter kildeskattelovens § 2, kan vælge beskatning efter grænsegængerreglerne i kildeskattelovens §§ 5 A-5 D, hvis betingelserne herfor er opfyldt. Efter disse regler behandles den skattepligtige ved indkomstopgørelsen i vidt omfang som en fuldt skattepligtig person.\n\n"
            "Beskatningen i Danmark omfatter dog fortsat kun den indkomst, som er skattepligtig til Danmark efter kildeskattelovens § 2."
        )
    return (
        "Med hjemmel i skatteforvaltningslovens § 26, stk. 1, har skatteministeren fastsat en kortere ligningsfrist for fysiske personer med enkle økonomiske forhold."
        + "\n\n"
        + first_paragraph
        + "\n\n"
        + second_paragraph
        + "\n\n"
        "Den korte ligningsfrist finder derfor ikke anvendelse, og du er dermed omfattet af den ordinære ligningsfrist i skatteforvaltningslovens § 26, stk. 1.\n\n"
        + deadline_lines
        + "\n\n"
        "Det fremgår af kildeskattelovens § 1, stk. 1, nr. 1, at personer, der har bopæl her i landet, er fuldt skattepligtige til Danmark. Ved afgørelsen af, om en person har bopæl i Danmark, lægges der blandt andet vægt på, om den pågældende faktisk har en bopælsmulighed i Danmark.\n\n"
        f"{bopæl_clause}, anser vi dig for at være fuldt skattepligtig til Danmark i {income_period_label} og omfattet af globalindkomstprincippet efter statsskattelovens § 4. Globalindkomstprincippet betyder, at alle indtægter er skattepligtige, uanset hvor de er optjent."
    )


def get_session_id(raw_session_id: str | None) -> str:
    session_id = (raw_session_id or "").strip()
    if not re.fullmatch(r"[a-zA-Z0-9_-]{8,80}", session_id):
        raise HTTPException(status_code=400, detail="Ugyldigt eller manglende chat_session_id")
    return session_id


def get_session_dir(session_id: str) -> Path:
    return CHAT_CONTEXT_DIR / session_id


def context_text_path(session_dir: Path, context_id: str) -> Path:
    return session_dir / f"{context_id}.txt"


def context_meta_path(session_dir: Path, context_id: str) -> Path:
    return session_dir / f"{context_id}.json"


def default_seed_marker_path(session_dir: Path) -> Path:
    return session_dir / ".default_seeded"


def list_chat_context_ids(session_id: str) -> list[str]:
    session_dir = get_session_dir(session_id)
    if not session_dir.exists():
        return []
    ids = [p.stem for p in session_dir.glob("*.json") if p.is_file()]
    ids.sort(
        key=lambda context_id: context_meta_path(session_dir, context_id).stat().st_mtime,
        reverse=True,
    )
    return ids


def load_context_meta(session_dir: Path, context_id: str) -> dict:
    meta_path = context_meta_path(session_dir, context_id)
    if not meta_path.exists():
        return {}
    try:
        return json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def ensure_default_guide_context_for_session(session_id: str) -> None:
    session_dir = get_session_dir(session_id)
    session_dir.mkdir(parents=True, exist_ok=True)
    marker_path = default_seed_marker_path(session_dir)
    if marker_path.exists():
        return

    if not DEFAULT_CHAT_GUIDE_PATH.exists():
        try:
            marker_path.write_text("missing_default_guide", encoding="utf-8")
        except Exception:
            pass
        return

    try:
        guide_text = normalize_text(DEFAULT_CHAT_GUIDE_PATH.read_text(encoding="utf-8"))
    except Exception:
        try:
            marker_path.write_text("default_guide_read_error", encoding="utf-8")
        except Exception:
            pass
        return

    if not guide_text:
        try:
            marker_path.write_text("default_guide_empty", encoding="utf-8")
        except Exception:
            pass
        return

    guide_text, was_truncated = truncate_text(guide_text, MAX_CHAT_CONTEXT_PER_FILE_CHARS)
    context_id = uuid4().hex
    context_text_path(session_dir, context_id).write_text(guide_text, encoding="utf-8")
    note = "Default skriveguide indlæst automatisk ved ny session"
    if was_truncated:
        note += ". Indhold blev afkortet."
    metadata = {
        "context_id": context_id,
        "filename": "Skriveguide.md",
        "file_type": "tekst",
        "size_chars": len(guide_text),
        "extraction_note": note,
    }
    context_meta_path(session_dir, context_id).write_text(
        json.dumps(metadata, ensure_ascii=False),
        encoding="utf-8",
    )
    try:
        marker_path.write_text("default_seeded", encoding="utf-8")
    except Exception:
        pass


def build_chat_context_list_response(session_id: str) -> ChatContextListResponse:
    ensure_default_guide_context_for_session(session_id)
    session_dir = get_session_dir(session_id)
    files: list[ChatContextFileResponse] = []
    for context_id in list_chat_context_ids(session_id):
        meta = load_context_meta(session_dir, context_id)
        filename = str(meta.get("filename", "ukendt_fil"))
        file_type = str(meta.get("file_type", "ukendt"))
        extraction_note = str(meta.get("extraction_note", "")).strip() or None
        size_chars = int(meta.get("size_chars", 0) or 0)
        files.append(
            ChatContextFileResponse(
                context_id=context_id,
                filename=filename,
                file_type=file_type,
                size_chars=size_chars,
                extraction_note=extraction_note,
            )
        )
    return ChatContextListResponse(files=files)


def load_chat_context_text(session_id: str) -> str:
    ensure_default_guide_context_for_session(session_id)
    session_dir = get_session_dir(session_id)
    context_ids = list_chat_context_ids(session_id)
    if not context_ids:
        return ""
    blocks: list[str] = []
    total_chars = 0
    for context_id in context_ids:
        meta = load_context_meta(session_dir, context_id)
        filename = str(meta.get("filename", "ukendt_fil"))
        file_type = str(meta.get("file_type", "ukendt"))
        try:
            text = context_text_path(session_dir, context_id).read_text(encoding="utf-8").strip()
        except Exception:
            continue
        if not text:
            continue
        block = f"[Kontekstfil: {filename} | type: {file_type}]\n{text}\n"
        remaining = MAX_CHAT_CONTEXT_CHARS - total_chars
        if remaining <= 0:
            break
        if len(block) > remaining:
            block = block[:remaining] + "\n[NOTE: Kontekst afkortet pga. længde]\n"
        blocks.append(block)
        total_chars += len(block)
        if total_chars >= MAX_CHAT_CONTEXT_CHARS:
            break
    return "\n".join(blocks).strip()


def ocr_image_bytes_with_openai(client: OpenAI, image_bytes: bytes, mime_type: str) -> str:
    image_b64 = base64.b64encode(image_bytes).decode("ascii")
    response = client.responses.create(
        model=PRIMARY_MODEL,
        input=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": (
                            "Du udfører OCR. Returner kun den udlæste tekst, "
                            "bevar linjeskift, og opfind ikke manglende tekst."
                        ),
                    },
                    {"type": "input_image", "image_url": f"data:{mime_type};base64,{image_b64}"},
                ],
            }
        ],
        reasoning={"effort": "medium"},
    )
    return normalize_text(str(getattr(response, "output_text", "") or ""))


def get_openai_client_for_ocr(existing_client: OpenAI | None) -> OpenAI:
    if existing_client is not None:
        return existing_client
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(
            status_code=500,
            detail="OPENAI_API_KEY mangler på server (kræves til OCR af billeder/scannede PDF'er)",
        )
    return OpenAI()


def extract_docx_text(data: bytes) -> tuple[str, str]:
    document = Document(io.BytesIO(data))
    lines: list[str] = []
    paragraph_count = 0
    for para in document.paragraphs:
        text = normalize_text(para.text)
        if text:
            lines.append(text)
            paragraph_count += 1
            if paragraph_count >= MAX_DOCX_PARAGRAPHS:
                break
    for table in document.tables:
        for row in table.rows:
            row_values = [normalize_text(cell.text) for cell in row.cells]
            row_values = [value for value in row_values if value]
            if row_values:
                lines.append(" | ".join(row_values))
    return "\n".join(lines).strip(), "DOCX udtrukket fra afsnit/tabeller"


def extract_excel_text(data: bytes, ext: str) -> tuple[str, str]:
    lines: list[str] = []
    sheets_processed = 0
    rows_processed = 0
    if ext == ".xlsx":
        workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        for sheet in workbook.worksheets[:MAX_EXCEL_SHEETS]:
            sheets_processed += 1
            lines.append(f"[Ark: {sheet.title}]")
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                if row_count >= MAX_EXCEL_ROWS_PER_SHEET:
                    lines.append("[NOTE: Flere rækker er afkortet]")
                    break
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if not values:
                    continue
                lines.append(" | ".join(values))
                row_count += 1
                rows_processed += 1
    else:
        workbook = xlrd.open_workbook(file_contents=data)
        for sheet in workbook.sheets()[:MAX_EXCEL_SHEETS]:
            sheets_processed += 1
            lines.append(f"[Ark: {sheet.name}]")
            row_count = 0
            for row_idx in range(sheet.nrows):
                if row_count >= MAX_EXCEL_ROWS_PER_SHEET:
                    lines.append("[NOTE: Flere rækker er afkortet]")
                    break
                row_values = [
                    str(sheet.cell_value(row_idx, col_idx)).strip()
                    for col_idx in range(sheet.ncols)
                    if str(sheet.cell_value(row_idx, col_idx)).strip()
                ]
                if not row_values:
                    continue
                lines.append(" | ".join(row_values))
                row_count += 1
                rows_processed += 1
    note = f"Excel udtrukket fra {sheets_processed} ark og {rows_processed} rækker"
    return "\n".join(lines).strip(), note


def extract_pdf_text_with_ocr(client: OpenAI | None, data: bytes) -> tuple[str, str]:
    pages: list[str] = []
    ocr_pages = 0
    processed_pages = 0

    try:
        reader = PdfReader(io.BytesIO(data))
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"PDF kunne ikke læses: {exc}") from exc

    page_count = min(len(reader.pages), doc.page_count, MAX_PDF_PAGES)
    for idx in range(page_count):
        processed_pages += 1
        page_text = normalize_text(reader.pages[idx].extract_text() or "")
        if len(page_text) >= MIN_PDF_PAGE_TEXT_FOR_NO_OCR:
            pages.append(f"[Side {idx + 1}]\n{page_text}")
            continue

        if ocr_pages >= MAX_PDF_OCR_PAGES:
            if page_text:
                pages.append(f"[Side {idx + 1}]\n{page_text}")
            continue

        pdf_page = doc.load_page(idx)
        pixmap = pdf_page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
        image_bytes = pixmap.tobytes("png")
        try:
            ocr_client = get_openai_client_for_ocr(client)
            ocr_text = ocr_image_bytes_with_openai(ocr_client, image_bytes, "image/png")
        except HTTPException:
            ocr_text = ""
        if ocr_text:
            ocr_pages += 1
            pages.append(f"[Side {idx + 1} OCR]\n{ocr_text}")
        elif page_text:
            pages.append(f"[Side {idx + 1}]\n{page_text}")

    note = f"PDF behandlet: {processed_pages} sider, OCR anvendt på {ocr_pages} sider"
    return "\n\n".join(pages).strip(), note


def extract_context_text(client: OpenAI | None, filename: str, ext: str, data: bytes) -> tuple[str, str, str]:
    if ext in {".md", ".txt"}:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise HTTPException(status_code=400, detail="Tekstfilen skal være UTF-8") from exc
        return normalize_text(text), "tekst", "Tekstfil indlæst"

    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
        }
        ocr_client = get_openai_client_for_ocr(client)
        text = ocr_image_bytes_with_openai(ocr_client, data, mime_map[ext])
        return text, "billede", "OCR udført via OpenAI vision"

    if ext == ".docx":
        text, note = extract_docx_text(data)
        return text, "word", note

    if ext in {".xlsx", ".xls"}:
        text, note = extract_excel_text(data, ext)
        return text, "excel", note

    if ext == ".pdf":
        text, note = extract_pdf_text_with_ocr(client, data)
        return text, "pdf", note

    raise HTTPException(status_code=400, detail="Filtype understøttes ikke")


@app.get("/api/chat/context", response_model=ChatContextListResponse)
def chat_context_list(
    x_chat_session_id: str | None = Header(default=None, alias="X-Chat-Session-Id"),
) -> ChatContextListResponse:
    session_id = get_session_id(x_chat_session_id)
    return build_chat_context_list_response(session_id)


@app.post("/api/chat/context", response_model=ChatContextListResponse)
async def upload_chat_context(
    file: UploadFile = File(...),
    x_chat_session_id: str | None = Header(default=None, alias="X-Chat-Session-Id"),
) -> ChatContextListResponse:
    session_id = get_session_id(x_chat_session_id)
    filename = file.filename or "context.txt"
    ext = Path(filename).suffix.lower()
    allowed_ext = {".md", ".txt", ".png", ".jpg", ".jpeg", ".webp", ".docx", ".pdf", ".xlsx", ".xls"}
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail="Filtype ikke tilladt")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Filen er tom")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="Filen er for stor")

    extracted_text, file_type, note = extract_context_text(None, filename, ext, data)
    if not extracted_text:
        raise HTTPException(status_code=400, detail="Kunne ikke udtrække brugbar tekst fra filen")

    extracted_text = normalize_text(extracted_text)
    extracted_text, was_truncated = truncate_text(extracted_text, MAX_CHAT_CONTEXT_PER_FILE_CHARS)
    if was_truncated:
        note = f"{note}. Indhold blev afkortet."

    CHAT_CONTEXT_DIR.mkdir(parents=True, exist_ok=True)
    session_dir = get_session_dir(session_id)
    session_dir.mkdir(parents=True, exist_ok=True)

    context_id = uuid4().hex
    context_text_path(session_dir, context_id).write_text(extracted_text, encoding="utf-8")
    metadata = {
        "context_id": context_id,
        "filename": filename,
        "file_type": file_type,
        "size_chars": len(extracted_text),
        "extraction_note": note,
    }
    context_meta_path(session_dir, context_id).write_text(
        json.dumps(metadata, ensure_ascii=False),
        encoding="utf-8",
    )
    return build_chat_context_list_response(session_id)


@app.delete("/api/chat/context/{context_id}", response_model=ChatContextListResponse)
def delete_chat_context(
    context_id: str,
    x_chat_session_id: str | None = Header(default=None, alias="X-Chat-Session-Id"),
) -> ChatContextListResponse:
    session_id = get_session_id(x_chat_session_id)
    if not re.fullmatch(r"[a-f0-9]{32}", context_id):
        raise HTTPException(status_code=400, detail="Ugyldigt context_id")
    session_dir = get_session_dir(session_id)
    text_file = context_text_path(session_dir, context_id)
    meta_file = context_meta_path(session_dir, context_id)
    if not text_file.exists() and not meta_file.exists():
        raise HTTPException(status_code=404, detail="Kontekstfil ikke fundet")
    try:
        if text_file.exists():
            text_file.unlink()
        if meta_file.exists():
            meta_file.unlink()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke fjerne kontekstfil: {exc}") from exc
    return build_chat_context_list_response(session_id)


@app.delete("/api/chat/context", response_model=ChatContextListResponse)
def clear_chat_context(
    x_chat_session_id: str | None = Header(default=None, alias="X-Chat-Session-Id"),
) -> ChatContextListResponse:
    session_id = get_session_id(x_chat_session_id)
    session_dir = get_session_dir(session_id)
    if session_dir.exists():
        for path in session_dir.glob("*"):
            if path.is_file():
                try:
                    path.unlink()
                except Exception:
                    continue
    return build_chat_context_list_response(session_id)


def _sse_line(data: dict | str) -> str:
    """Formatér dict som SSE data-linje."""
    import json as _json
    return f"data: {_json.dumps(data, ensure_ascii=False)}\n\n"


@app.post("/api/analyze", response_model=AnalyzeResponse)
def analyze(
    payload: AnalyzeRequest,
    stream: bool = Query(False, description="Stream svar som SSE"),
) -> AnalyzeResponse | StreamingResponse:
    question = payload.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="Spørgsmål må ikke være tomt")

    source_tab = (payload.source_tab or "").strip().lower()
    subtab = (payload.subtab or "").strip().lower()
    case_entry: dict | None = None
    case_store = get_case_store()
    case_id = str(payload.case_id or "").strip()
    case_user = str(payload.case_user or "").strip()
    if source_tab == "sagsbehandling" and case_id:
        if not case_user:
            raise HTTPException(status_code=400, detail="Bruger mangler for aktiv sag")
        case_entry = case_store.get_case(case_user, case_id)
        if not case_entry:
            raise HTTPException(status_code=404, detail="Sag ikke fundet")
    if source_tab == "sagsbehandling" and subtab == "skattepligt_ligningsfrist":
        try:
            answer = build_standard_text_ligningsfrist(payload.case_facts)
            parsed = {
                "output_text": answer,
                "citations": [],
                "retrieved_chunks": [],
                "used_vector_store_ids": [],
            }
            used_model = "regelmotor-ligningsfrist-v1"
            response_id = f"rule_{uuid4().hex}"
            log_question = (
                "Regelmotor input\n"
                f"- source_tab: {source_tab}\n"
                f"- subtab: {subtab}\n"
                f"- case_facts: {json.dumps(payload.case_facts or {}, ensure_ascii=False)}"
            )
            log_path = save_pdf_log(log_question, parsed, used_model)
            if case_entry and case_id and case_user:
                shared = dict(case_entry.get("shared_facts") or {})
                facts = payload.case_facts or {}
                if isinstance(facts, dict):
                    if facts.get("income_years"):
                        shared["income_years"] = str(facts.get("income_years"))
                    if facts.get("selected_trigger"):
                        shared["selected_trigger"] = str(facts.get("selected_trigger"))
                    if facts.get("residence_fact"):
                        shared["residence_fact"] = str(facts.get("residence_fact"))
                case_store.update_case(
                    username=case_user,
                    case_id=case_id,
                    patch={
                        "active_subtab": subtab,
                        "shared_facts": shared,
                        "subtab_outputs": {
                            subtab: {
                                "answer": answer,
                                "used_model": used_model,
                                "response_id": response_id,
                            }
                        },
                    },
                )
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Regelmotor fejlede: {exc}") from exc

        log_filename = log_path.name
        return AnalyzeResponse(
            answer=answer,
            used_model=used_model,
            response_id=response_id,
            citations=[],
            retrieval_results=[],
            log_pdf_filename=log_filename,
            log_pdf_url=f"/api/logs/{log_filename}",
        )

    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY mangler på server")

    vector_store_ids_override: list[str] | None = None
    models_override: list[str] | None = None
    instructions_override: str | None = None
    llm_question = question
    use_file_search = True
    if source_tab == "sagsbehandling" and subtab in SAGSBEHANDLING_VECTOR_STORES:
        vector_store_ids_override = SAGSBEHANDLING_VECTOR_STORES[subtab]
        models_override = SAGSBEHANDLING_MODELS.get(subtab)
        instructions_override = SAGSBEHANDLING_PROMPTS.get(subtab)
        facts_block = format_case_facts_for_llm(subtab, payload.case_facts)
        if facts_block:
            llm_question = (
                question
                + "\n\n---\n"
                + facts_block
                + "\n---\n"
                + "Brug disse faktiske oplysninger sammen med de fundne kilder i file_search."
            )
        case_block = format_case_shared_facts_for_llm(case_entry, subtab)
        if case_block:
            llm_question = (
                llm_question
                + "\n\n---\n"
                + case_block
                + "\n---\n"
                + "Brug de fælles sagsfakta som baggrund, men følg file_search-kilder ved konflikt."
            )
        decision_block = format_sags_decision_package_for_llm(payload.sags_decision_package)
        if decision_block:
            llm_question = (
                llm_question
                + "\n\n---\n"
                + decision_block
                + "\n---\n"
                + "Brug beslutningspakken som primær struktur for analysen. "
                + "Skeln tydeligt mellem konstaterede fakta, afledte præmisser, metode, foreløbig beskatningsret, konflikter og advarsler."
            )

    context_log_ids_to_use: list[str] = []
    if payload.context_log_ids:
        seen_ids: set[str] = set()
        for item in payload.context_log_ids:
            clean_id = str(item or "").strip()
            if not clean_id or clean_id in seen_ids:
                continue
            seen_ids.add(clean_id)
            context_log_ids_to_use.append(clean_id)
    elif payload.context_log_id and str(payload.context_log_id or "").strip():
        context_log_ids_to_use = [str(payload.context_log_id or "").strip()]

    if source_tab == "sagsbehandling" and context_log_ids_to_use:
        if len(context_log_ids_to_use) > MAX_SAGS_CONTEXT_LOGS:
            raise HTTPException(
                status_code=400,
                detail=f"Du kan højst vælge {MAX_SAGS_CONTEXT_LOGS} analyse-kontekster ad gangen",
            )
        if not payload.context_approved:
            raise HTTPException(
                status_code=400,
                detail="Analyse-kontekst skal godkendes i UI før den kan bruges",
            )
        if not (payload.context_user or "").strip():
            raise HTTPException(
                status_code=400,
                detail="Bruger mangler for valgt analyse-kontekst",
            )
        context_block = build_sags_context_from_analyse_logs(
            username=payload.context_user or "",
            context_log_ids=context_log_ids_to_use,
        )
        if context_block:
            llm_question = (
                llm_question
                + "\n\n---\n"
                + context_block
                + "\n---\n"
                + "Brug konteksten som baggrund. Hvis den strider mod file_search-kilder, følg file_search-kilderne."
            )

    legal_context_blocks_raw = payload.legal_context_blocks or []
    legal_context_blocks_to_use: list[str] = []
    for block in legal_context_blocks_raw:
        clean_block = str(block or "").strip()
        if not clean_block:
            continue
        clean_block, _ = truncate_text(clean_block, MAX_ANALYSE_LEGAL_CONTEXT_CHARS)
        legal_context_blocks_to_use.append(clean_block)
        if len(legal_context_blocks_to_use) >= MAX_ANALYSE_LEGAL_CONTEXT_BLOCKS:
            break
    if source_tab == "analyse" and legal_context_blocks_to_use:
        context_text = "\n\n".join(
            f"[Retskildekontekst {index + 1}]\n{block}"
            for index, block in enumerate(legal_context_blocks_to_use)
        )
        llm_question = (
            llm_question
            + "\n\n---\n"
            + context_text
            + "\n---\n"
            + "Brug denne retskildekontekst som primær baggrund for analysen."
        )
        if not payload.use_semantic_search_with_legal_context:
            use_file_search = False

    if stream:
        def gen():
            try:
                client = OpenAI()
                reasoning_effort = REASONING_EFFORT_LIGNINGSFRIST if instructions_override else None
                prompt_cache_key = PROMPT_CACHE_KEY_LIGNINGSFRIST if instructions_override else None
                for evt in analyze_question_stream(
                    client=client,
                    question=llm_question,
                    log_question=llm_question,
                    previous_response_id=payload.previous_response_id,
                    vector_store_ids=vector_store_ids_override,
                    instructions=instructions_override,
                    models_to_try=models_override,
                    reasoning_effort=reasoning_effort,
                    prompt_cache_key=prompt_cache_key,
                    use_file_search=use_file_search,
                ):
                    yield _sse_line(evt)
            except Exception as exc:
                yield _sse_line({"type": "error", "detail": str(exc)})

        return StreamingResponse(
            gen(),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    try:
        client = OpenAI()
        reasoning_effort = REASONING_EFFORT_LIGNINGSFRIST if instructions_override else None
        prompt_cache_key = PROMPT_CACHE_KEY_LIGNINGSFRIST if instructions_override else None
        parsed, used_model, response_id = analyze_question(
            client=client,
            question=llm_question,
            previous_response_id=payload.previous_response_id,
            vector_store_ids=vector_store_ids_override,
            instructions=instructions_override,
            models_to_try=models_override,
            reasoning_effort=reasoning_effort,
            prompt_cache_key=prompt_cache_key,
            use_file_search=use_file_search,
        )
        log_path = save_pdf_log(llm_question, parsed, used_model)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Analyse fejlede: {exc}") from exc

    log_filename = log_path.name
    if (
        source_tab == "sagsbehandling"
        and case_entry
        and case_id
        and case_user
    ):
        case_store.update_case(
            username=case_user,
            case_id=case_id,
            patch={
                "active_subtab": subtab,
                "subtab_outputs": {
                    subtab: {
                        "answer": parsed.get("output_text", ""),
                        "used_model": used_model,
                        "response_id": response_id,
                    }
                },
            },
        )
    return AnalyzeResponse(
        answer=parsed.get("output_text", ""),
        used_model=used_model,
        response_id=response_id,
        citations=parsed.get("citations", []),
        retrieval_results=parsed.get("retrieved_chunks", []),
        log_pdf_filename=log_filename,
        log_pdf_url=f"/api/logs/{log_filename}",
    )


@app.get("/api/sagsbehandling/legal-basis", response_model=SagsLegalBasisResponse)
def sagsbehandling_legal_basis(subtab: str) -> SagsLegalBasisResponse:
    subtab_key = (subtab or "").strip().lower()
    vector_store_ids = SAGSBEHANDLING_VECTOR_STORES.get(subtab_key, [])
    vector_store_id = vector_store_ids[0] if vector_store_ids else None
    if not vector_store_id:
        return SagsLegalBasisResponse(subtab=subtab_key, vector_store_id=None, documents=[])
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY mangler på server")

    try:
        client = OpenAI()
        documents = list_vector_store_document_names(client, vector_store_id)
        return SagsLegalBasisResponse(
            subtab=subtab_key,
            vector_store_id=vector_store_id,
            documents=documents,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke hente retsgrundlag: {exc}") from exc


@app.get("/api/legal-sources/catalog", response_model=LegalSourcesCatalogResponse)
def get_legal_sources_catalog() -> LegalSourcesCatalogResponse:
    catalog_files = sorted(LEGAL_SOURCES_DIR.glob("*_catalog.json"))
    if not catalog_files:
        return LegalSourcesCatalogResponse(categories=[], documents=[])
    categories_by_id: dict[str, dict[str, object]] = {}
    documents: list[dict[str, object]] = []
    try:
        for catalog_path in catalog_files:
            payload = json.loads(catalog_path.read_text(encoding="utf-8"))
            categories = payload.get("categories")
            for category in categories if isinstance(categories, list) else []:
                category_id = str((category or {}).get("id", "")).strip().lower()
                if category_id and category_id not in categories_by_id:
                    categories_by_id[category_id] = category
            source_documents = payload.get("documents")
            if isinstance(source_documents, list):
                documents.extend(source_documents)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke læse retskildekatalog: {exc}") from exc
    return LegalSourcesCatalogResponse(
        categories=list(categories_by_id.values()),
        documents=documents,
    )


@app.get("/api/legal-sources/file/{source_id}")
def get_legal_source_file(source_id: str) -> FileResponse:
    file_path = resolve_legal_pdf_by_source_id(source_id)
    return FileResponse(path=file_path, media_type="application/pdf", filename=file_path.name)


@app.get("/api/legal-sources/section/{source_id}", response_model=LegalSourceSectionResponse)
def get_legal_source_section(
    source_id: str,
    page: int = Query(default=1, ge=1, description="1-indexed preview page"),
    chunk_size: int = Query(default=8, ge=1, le=100, description="Number of PDF pages per preview chunk"),
) -> LegalSourceSectionResponse:
    file_path = resolve_legal_pdf_by_source_id(source_id)
    cache_key, namespace = parse_legal_source_namespace(source_id)
    precomputed = load_precomputed_legal_previews(namespace).get(cache_key)
    if isinstance(precomputed, dict):
        pages = precomputed.get("pages")
        title = str(precomputed.get("title", file_path.stem))
        if isinstance(pages, list) and pages:
            total_pages = max(1, (len(pages) + chunk_size - 1) // chunk_size)
            safe_page = max(1, min(page, total_pages))
            start_idx = (safe_page - 1) * chunk_size
            end_idx = min(len(pages), start_idx + chunk_size)
            text_block = "\n\n".join(str(pages[idx] or "") for idx in range(start_idx, end_idx)).strip()
            return LegalSourceSectionResponse(
                source_id=source_id,
                title=title,
                text=clean_legal_preview_text(text_block, namespace=namespace),
                truncated=False,
                page=safe_page,
                total_pages=total_pages,
            )

    try:
        file_mtime = float(file_path.stat().st_mtime)
    except Exception:
        file_mtime = 0.0
    cached = LEGAL_SOURCE_PREVIEW_CACHE.get(cache_key)
    if (
        isinstance(cached, dict)
        and float(cached.get("mtime", -1.0)) == file_mtime
        and isinstance(cached.get("pages"), list)
    ):
        cached_pages = [str(item or "") for item in cached.get("pages", []) if str(item or "").strip()]
        if not cached_pages:
            cached_pages = [str(cached.get("text", ""))]
        total_pages = max(1, (len(cached_pages) + chunk_size - 1) // chunk_size)
        safe_page = max(1, min(page, total_pages))
        start_idx = (safe_page - 1) * chunk_size
        end_idx = min(len(cached_pages), start_idx + chunk_size)
        text_block = "\n\n".join(str(cached_pages[idx] or "") for idx in range(start_idx, end_idx)).strip()
        return LegalSourceSectionResponse(
            source_id=source_id,
            title=str(cached.get("title", file_path.stem)),
            text=clean_legal_preview_text(text_block, namespace=namespace),
            truncated=bool(cached.get("truncated", False)),
            page=safe_page,
            total_pages=total_pages,
        )

    preview_pages = extract_pdf_preview_pages(file_path, namespace=namespace)
    if not preview_pages:
        preview_text, truncated = extract_pdf_preview_text(file_path, namespace=namespace)
        preview_pages = [preview_text]
    else:
        truncated = False
    LEGAL_SOURCE_PREVIEW_CACHE[cache_key] = {
        "mtime": file_mtime,
        "title": file_path.stem,
        "pages": preview_pages,
        "text": preview_pages[0] if preview_pages else "",
        "truncated": truncated,
    }
    total_pages = max(1, (len(preview_pages) + chunk_size - 1) // chunk_size)
    safe_page = max(1, min(page, total_pages))
    start_idx = (safe_page - 1) * chunk_size
    end_idx = min(len(preview_pages), start_idx + chunk_size)
    text_block = "\n\n".join(str(preview_pages[idx] or "") for idx in range(start_idx, end_idx)).strip()
    return LegalSourceSectionResponse(
        source_id=source_id,
        title=file_path.stem,
        text=clean_legal_preview_text(text_block, namespace=namespace),
        truncated=truncated,
        page=safe_page,
        total_pages=total_pages,
    )


@app.post("/api/cases", response_model=CaseGetResponse)
def create_case_endpoint(payload: CaseCreateRequest) -> CaseGetResponse:
    """Start ny sag for bruger."""
    try:
        entry = get_case_store().create_case(payload.user, payload.title)
        return CaseGetResponse(
            id=entry.get("id", ""),
            title=entry.get("title", "Ny sag"),
            status=entry.get("status", "open"),
            created_at=entry.get("created_at", ""),
            updated_at=entry.get("updated_at", ""),
            active_subtab=entry.get("active_subtab", "skattepligt_ligningsfrist"),
            shared_facts=entry.get("shared_facts", {}),
            subtab_outputs=entry.get("subtab_outputs", {}),
            locked_by_subtab=entry.get("locked_by_subtab", {}),
            facts_locked_by_subtab=entry.get("facts_locked_by_subtab", {}),
            facts_by_subtab=entry.get("facts_by_subtab", {}),
            context_by_subtab=entry.get("context_by_subtab", {}),
            messages_by_subtab=entry.get("messages_by_subtab", {}),
            previous_response_id_by_subtab=entry.get("previous_response_id_by_subtab", {}),
            used_model_by_subtab=entry.get("used_model_by_subtab", {}),
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke oprette sag: {exc}") from exc


@app.get("/api/cases", response_model=CaseListResponse)
def list_cases_endpoint(user: str = Query(..., min_length=1)) -> CaseListResponse:
    """Liste af sager for bruger."""
    entries = get_case_store().list_cases(user)
    return CaseListResponse(entries=entries)


@app.get("/api/cases/{case_id}", response_model=CaseGetResponse)
def get_case_endpoint(case_id: str, user: str = Query(..., min_length=1)) -> CaseGetResponse:
    """Hent fuld sag."""
    entry = get_case_store().get_case(user, case_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Sag ikke fundet")
    return CaseGetResponse(
        id=entry.get("id", ""),
        title=entry.get("title", "Ny sag"),
        status=entry.get("status", "open"),
        created_at=entry.get("created_at", ""),
        updated_at=entry.get("updated_at", ""),
        active_subtab=entry.get("active_subtab", "skattepligt_ligningsfrist"),
        shared_facts=entry.get("shared_facts", {}),
        subtab_outputs=entry.get("subtab_outputs", {}),
        locked_by_subtab=entry.get("locked_by_subtab", {}),
        facts_locked_by_subtab=entry.get("facts_locked_by_subtab", {}),
        facts_by_subtab=entry.get("facts_by_subtab", {}),
        context_by_subtab=entry.get("context_by_subtab", {}),
        messages_by_subtab=entry.get("messages_by_subtab", {}),
        previous_response_id_by_subtab=entry.get("previous_response_id_by_subtab", {}),
        used_model_by_subtab=entry.get("used_model_by_subtab", {}),
    )


@app.patch("/api/cases/{case_id}", response_model=CaseGetResponse)
def update_case_endpoint(case_id: str, payload: CaseUpdateRequest) -> CaseGetResponse:
    """Opdater sag (delvis update)."""
    patch: dict[str, object] = {}
    if payload.title is not None:
        patch["title"] = payload.title
    if payload.status is not None:
        patch["status"] = payload.status
    if payload.active_subtab is not None:
        patch["active_subtab"] = payload.active_subtab
    if payload.shared_facts is not None:
        patch["shared_facts"] = payload.shared_facts
    if payload.subtab_outputs is not None:
        patch["subtab_outputs"] = payload.subtab_outputs
    if payload.locked_by_subtab is not None:
        patch["locked_by_subtab"] = payload.locked_by_subtab
    if payload.facts_locked_by_subtab is not None:
        patch["facts_locked_by_subtab"] = payload.facts_locked_by_subtab
    if payload.facts_by_subtab is not None:
        patch["facts_by_subtab"] = payload.facts_by_subtab
    if payload.context_by_subtab is not None:
        patch["context_by_subtab"] = payload.context_by_subtab
    if payload.messages_by_subtab is not None:
        patch["messages_by_subtab"] = {
            key: [{"role": msg.role, "text": msg.text} for msg in messages]
            for key, messages in payload.messages_by_subtab.items()
        }
    if payload.previous_response_id_by_subtab is not None:
        patch["previous_response_id_by_subtab"] = payload.previous_response_id_by_subtab
    if payload.used_model_by_subtab is not None:
        patch["used_model_by_subtab"] = payload.used_model_by_subtab
    entry = get_case_store().update_case(payload.user, case_id, patch)
    if not entry:
        raise HTTPException(status_code=404, detail="Sag ikke fundet")
    return CaseGetResponse(
        id=entry.get("id", ""),
        title=entry.get("title", "Ny sag"),
        status=entry.get("status", "open"),
        created_at=entry.get("created_at", ""),
        updated_at=entry.get("updated_at", ""),
        active_subtab=entry.get("active_subtab", "skattepligt_ligningsfrist"),
        shared_facts=entry.get("shared_facts", {}),
        subtab_outputs=entry.get("subtab_outputs", {}),
        locked_by_subtab=entry.get("locked_by_subtab", {}),
        facts_locked_by_subtab=entry.get("facts_locked_by_subtab", {}),
        facts_by_subtab=entry.get("facts_by_subtab", {}),
        context_by_subtab=entry.get("context_by_subtab", {}),
        messages_by_subtab=entry.get("messages_by_subtab", {}),
        previous_response_id_by_subtab=entry.get("previous_response_id_by_subtab", {}),
        used_model_by_subtab=entry.get("used_model_by_subtab", {}),
    )


@app.delete("/api/cases/{case_id}", response_model=CaseListResponse)
def delete_case_endpoint(case_id: str, user: str = Query(..., min_length=1)) -> CaseListResponse:
    """Slet sag."""
    deleted = get_case_store().delete_case(user, case_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Sag ikke fundet")
    entries = get_case_store().list_cases(user)
    return CaseListResponse(entries=entries)


@app.post("/api/chat", response_model=ChatResponse)
def chat(
    payload: ChatRequest,
    x_chat_session_id: str | None = Header(default=None, alias="X-Chat-Session-Id"),
    stream: bool = Query(False, description="Stream svar som SSE"),
) -> ChatResponse | StreamingResponse:
    session_id = get_session_id(x_chat_session_id)
    message = payload.message.strip()
    if not message:
        raise HTTPException(status_code=400, detail="Chatbesked må ikke være tom")
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY mangler på server")

    try:
        client = OpenAI()
        context_text = load_chat_context_text(session_id)
        chat_instructions = CHAT_INSTRUCTIONS
        if context_text:
            chat_instructions = (
                CHAT_INSTRUCTIONS
                + "\n\nYderligere uploadet kontekst til chat (skal anvendes i samspil med ovenstående):\n"
                + context_text
            )

        if stream:
            def chat_gen():
                try:
                    req = {
                        "model": PRIMARY_MODEL,
                        "instructions": chat_instructions,
                        "input": message,
                        "reasoning": {"effort": REASONING_EFFORT_CHAT},
                        "prompt_cache_key": PROMPT_CACHE_KEY_CHAT,
                        "prompt_cache_retention": PROMPT_CACHE_RETENTION,
                        "stream": True,
                    }
                    if payload.previous_response_id:
                        req["previous_response_id"] = payload.previous_response_id
                    t0 = time.perf_counter()
                    stream_resp = client.responses.create(**req)
                    for event in stream_resp:
                        ev_type = getattr(event, "type", None) or (event.get("type") if isinstance(event, dict) else None)
                        if ev_type == "response.output_text.delta":
                            delta = getattr(event, "delta", None) or (event.get("delta", "") if isinstance(event, dict) else "")
                            if delta:
                                yield _sse_line({"type": "delta", "text": delta})
                        elif ev_type == "response.completed":
                            resp_obj = getattr(event, "response", None) or (event.get("response") if isinstance(event, dict) else None)
                            duration_ms = (time.perf_counter() - t0) * 1000
                            answer = str(getattr(resp_obj, "output_text", "") or "") if resp_obj else ""
                            response_id = str(getattr(resp_obj, "id", "") or "") if resp_obj else ""
                            usage = getattr(resp_obj, "usage", None) if resp_obj else None
                            inp_tok = getattr(usage, "input_tokens", 0) if usage else 0
                            out_tok = getattr(usage, "output_tokens", 0) if usage else 0
                            req_id = getattr(resp_obj, "_request_id", None) if resp_obj else None
                            _log.info(
                                "perf flow=chat model=%s duration_ms=%.0f x_request_id=%s input_tokens=%s output_tokens=%s",
                                PRIMARY_MODEL, duration_ms, req_id or "?", inp_tok, out_tok,
                            )
                            yield _sse_line({
                                "type": "done",
                                "answer": answer.strip(),
                                "used_model": PRIMARY_MODEL,
                                "response_id": response_id,
                            })
                except Exception as exc:
                    yield _sse_line({"type": "error", "detail": str(exc)})

            return StreamingResponse(
                chat_gen(),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
            )

        request_payload: dict[str, object] = {
            "model": PRIMARY_MODEL,
            "instructions": chat_instructions,
            "input": message,
            "reasoning": {"effort": REASONING_EFFORT_CHAT},
            "prompt_cache_key": PROMPT_CACHE_KEY_CHAT,
            "prompt_cache_retention": PROMPT_CACHE_RETENTION,
        }
        if payload.previous_response_id:
            request_payload["previous_response_id"] = payload.previous_response_id
        t0 = time.perf_counter()
        resp = client.responses.create(**request_payload)
        duration_ms = (time.perf_counter() - t0) * 1000
        answer = str(getattr(resp, "output_text", "") or "").strip()
        response_id = str(getattr(resp, "id", "") or "")
        usage = getattr(resp, "usage", None)
        input_tokens = getattr(usage, "input_tokens", 0) if usage else 0
        output_tokens = getattr(usage, "output_tokens", 0) if usage else 0
        req_id = getattr(resp, "_request_id", None)
        _log.info(
            "perf flow=chat model=%s duration_ms=%.0f x_request_id=%s input_tokens=%s output_tokens=%s reasoning_effort=%s",
            PRIMARY_MODEL,
            duration_ms,
            req_id or "?",
            input_tokens,
            output_tokens,
            REASONING_EFFORT_CHAT,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Chat fejlede: {exc}") from exc

    return ChatResponse(
        answer=answer or "Intet svar returneret.",
        used_model=PRIMARY_MODEL,
        response_id=response_id,
    )


@app.post("/api/chat/export-pdf", response_model=ChatExportResponse)
def export_chat_pdf(
    payload: ChatExportRequest,
    x_chat_session_id: str | None = Header(default=None, alias="X-Chat-Session-Id"),
) -> ChatExportResponse:
    get_session_id(x_chat_session_id)
    messages = [
        {"role": msg.role.strip(), "text": msg.text.strip()}
        for msg in payload.messages
        if msg.text.strip()
    ]
    if not messages:
        raise HTTPException(status_code=400, detail="Der er ingen chatbeskeder at gemme")

    try:
        log_path = save_chat_pdf_log(messages, PRIMARY_MODEL)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke gemme chat-PDF: {exc}") from exc

    log_filename = log_path.name
    return ChatExportResponse(
        log_pdf_filename=log_filename,
        log_pdf_url=f"/api/logs/{log_filename}",
    )


@app.post("/api/analyse-logs", response_model=AnalyseLogSaveResponse)
def save_analyse_log_endpoint(payload: AnalyseLogSaveRequest) -> AnalyseLogSaveResponse:
    """Gem analyse-log med LLM-genereret titel."""
    try:
        result = save_analyse_log(
            username=payload.user,
            session_id=payload.session_id,
            question=payload.question,
            answer=payload.answer,
            citations=payload.citations,
            retrieval_results=payload.retrieval_results,
            used_model=payload.used_model,
            log_question=payload.log_question,
            used_vector_store_ids=payload.used_vector_store_ids,
            log_pdf_filename=payload.log_pdf_filename,
            log_pdf_url=payload.log_pdf_url,
            messages=[{"role": msg.role, "text": msg.text} for msg in payload.messages],
            last_response_id=payload.last_response_id,
        )
        return AnalyseLogSaveResponse(**result)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke gemme log: {exc}") from exc


@app.get("/api/analyse-logs", response_model=AnalyseLogListResponse)
def list_analyse_logs_endpoint(user: str = Query(..., min_length=1)) -> AnalyseLogListResponse:
    """Liste gemte analyse-logs for bruger."""
    entries = list_analyse_logs(user)
    return AnalyseLogListResponse(entries=entries)


@app.get("/api/analyse-logs/{entry_id}", response_model=AnalyseLogGetResponse)
def get_analyse_log_endpoint(
    entry_id: str,
    user: str = Query(..., min_length=1),
) -> AnalyseLogGetResponse:
    """Hent fuld analyse-log efter id."""
    entry = get_analyse_log(user, entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Log ikke fundet")
    return AnalyseLogGetResponse(
        id=entry["id"],
        session_id=entry.get("session_id"),
        created_at=entry["created_at"],
        title=entry.get("title", "Uden titel"),
        question=entry.get("question", ""),
        answer=entry.get("answer", ""),
        citations=entry.get("citations", []),
        retrieval_results=entry.get("retrieval_results", []),
        used_model=entry.get("used_model", ""),
        used_vector_store_ids=entry.get("used_vector_store_ids", []),
        log_pdf_filename=entry.get("log_pdf_filename"),
        log_pdf_url=entry.get("log_pdf_url"),
        messages=[
            {"role": str(msg.get("role", "")).strip(), "text": str(msg.get("text", "")).strip()}
            for msg in (entry.get("messages") or [])
            if str(msg.get("text", "")).strip()
        ],
        last_response_id=entry.get("last_response_id"),
    )


@app.delete("/api/analyse-logs/{entry_id}", response_model=AnalyseLogListResponse)
def delete_analyse_log_endpoint(
    entry_id: str,
    user: str = Query(..., min_length=1),
) -> AnalyseLogListResponse:
    """Slet analyse-log efter id."""
    deleted = delete_analyse_log(user, entry_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Log ikke fundet")
    entries = list_analyse_logs(user)
    return AnalyseLogListResponse(entries=entries)


@app.post("/api/chat-logs", response_model=ChatLogSaveResponse)
def save_chat_log_endpoint(payload: ChatLogSaveRequest) -> ChatLogSaveResponse:
    """Gem eller opdater chat-log for en session."""
    try:
        result = save_chat_log(
            username=payload.user,
            session_id=payload.session_id,
            messages=[{"role": msg.role, "text": msg.text} for msg in payload.messages],
            used_model=payload.used_model,
            last_response_id=payload.last_response_id,
        )
        return ChatLogSaveResponse(**result)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Kunne ikke gemme chat-log: {exc}") from exc


@app.get("/api/chat-logs", response_model=ChatLogListResponse)
def list_chat_logs_endpoint(user: str = Query(..., min_length=1)) -> ChatLogListResponse:
    """Liste gemte chat-logs for bruger."""
    entries = list_chat_logs(user)
    return ChatLogListResponse(entries=entries)


@app.get("/api/chat-logs/{entry_id}", response_model=ChatLogGetResponse)
def get_chat_log_endpoint(
    entry_id: str,
    user: str = Query(..., min_length=1),
) -> ChatLogGetResponse:
    """Hent fuld chat-log efter id."""
    entry = get_chat_log(user, entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Chat-log ikke fundet")
    return ChatLogGetResponse(
        id=entry.get("id", ""),
        session_id=entry.get("session_id", ""),
        title=entry.get("title", "Chat uden titel"),
        created_at=entry.get("created_at", ""),
        updated_at=entry.get("updated_at", entry.get("created_at", "")),
        used_model=entry.get("used_model", ""),
        last_response_id=entry.get("last_response_id"),
        messages=[
            {"role": str(msg.get("role", "")).strip(), "text": str(msg.get("text", "")).strip()}
            for msg in (entry.get("messages") or [])
            if str(msg.get("text", "")).strip()
        ],
    )


@app.delete("/api/chat-logs/{entry_id}", response_model=ChatLogListResponse)
def delete_chat_log_endpoint(
    entry_id: str,
    user: str = Query(..., min_length=1),
) -> ChatLogListResponse:
    """Slet chat-log efter id."""
    deleted = delete_chat_log(user, entry_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Chat-log ikke fundet")
    entries = list_chat_logs(user)
    return ChatLogListResponse(entries=entries)


@app.get("/api/logs/{filename}")
def get_log_file(filename: str) -> FileResponse:
    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Kun PDF logs er tilladt")

    candidate = (LOG_DIR / filename).resolve()
    log_dir_resolved = LOG_DIR.resolve()

    # Avoid path traversal and ensure file stays in logs directory.
    if log_dir_resolved not in candidate.parents or not candidate.exists():
        raise HTTPException(status_code=404, detail="Logfil ikke fundet")

    return FileResponse(path=candidate, media_type="application/pdf", filename=filename)
